# src/templify/core/templifier.py
from docx import Document
import re

class Templifier:
    """
    Applies a config (titles_config + docx_config) to new plaintext
    to generate a styled DOCX document.
    """

    def __init__(self, titles_config: dict, docx_config: dict):
        self.titles_config = titles_config
        self.docx_config = docx_config

    def build_docx(self, plaintext_lines: list[str], output_path: str):
        doc = Document()

        style_map = self._build_style_map()

        for line in plaintext_lines:
            matched_section = self._match_section(line)

            if matched_section:
                style_info = style_map.get(matched_section, {})
                para = doc.add_paragraph(line)
                self._apply_style(para, style_info)
            else:
                doc.add_paragraph(line)  # fallback

        doc.save(output_path)
        return output_path

    def _build_style_map(self):
        mapping = {}
        for group in self.docx_config.get("layout_groups", []):
            for section in group.get("section_types", []):
                mapping[section["section_type"]] = {
                    "font": section.get("font", {}),
                    "paragraph": section.get("paragraph", {}),
                }
        return mapping

    def _match_section(self, line: str) -> str | None:
        for title in self.titles_config.get("titles", []):
            pattern = title["title_detection"]["pattern"]
            flags = 0 if title["title_detection"]["case_sensitive"] else re.IGNORECASE
            if re.fullmatch(pattern, line, flags):
                return title["section_type"]
        return None

    def _apply_style(self, para, style_info):
        font = para.runs[0].font if para.runs else para.add_run("").font
        f = style_info.get("font", {})

        if "name" in f: font.name = f["name"]
        if "size" in f: font.size = f["size"]
        if f.get("bold"): font.bold = True
        if f.get("italic"): font.italic = True
        if f.get("underline"): font.underline = True

        # TODO: apply paragraph spacing / alignment
