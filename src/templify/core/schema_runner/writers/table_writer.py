# src/templify/core/schema_runner/writers/table_writer.py

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class TableWriter:
    """
    Handles writing tables into the document.
    Expects schema format:
    {
      "id": "tbl_1",
      "style": {"style_id": "TableGrid", "banded": true, "autofit": true},
      "rows": [
        {"cells": [
          {"text": "Project", "style": {"font": {"bold": true}, "paragraph": {"alignment": "center"}}},
          {"text": "Role", "style": {}},
          {"text": "Year", "style": {}}
        ]},
        {"cells": [
          {"text": "Templify", "style": {}},
          {"text": "Developer", "style": {}},
          {"text": "2025", "style": {}}
        ]}
      ]
    }
    """

    def __init__(self, document):
        self.doc = document

    def write(self, descriptor: dict, style: dict = None):
        """
        Write a table using a descriptor from schema.
        :param descriptor: dict from schema["tables"][i]
        :param style: optional override style dict (table-level)
        """
        rows = descriptor.get("rows", [])
        if not rows:
            return None

        nrows = len(rows)
        ncols = len(rows[0]["cells"])

        table = self.doc.add_table(rows=nrows, cols=ncols)

        # Apply table style
        table_style = descriptor.get("style", {})
        if style:
            table_style.update(style)
        if "style_id" in table_style:
            try:
                table.style = table_style["style_id"]
            except KeyError:
                pass
        if "autofit" in table_style:
            table.autofit = table_style["autofit"]

        # Fill cells
        for i, row in enumerate(rows):
            for j, cell in enumerate(row["cells"]):
                cell_obj = table.cell(i, j)
                text = cell.get("text", "")
                cell_obj.text = text

                # Apply cell-level style
                if "style" in cell:
                    self._apply_cell_style(cell_obj, cell["style"])

        return table

    def _apply_cell_style(self, cell, style: dict):
        """
        Apply font + paragraph styles to a cell.
        """
        if not style:
            return

        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = p.runs[0] if p.runs else p.add_run()

        # Font overrides
        if "font" in style:
            font = style["font"]
            if "name" in font:
                run.font.name = font["name"]
            if "size" in font and font["size"]:
                run.font.size = Pt(font["size"])
            if "bold" in font:
                run.font.bold = font["bold"]
            if "italic" in font:
                run.font.italic = font["italic"]
            if "underline" in font:
                run.font.underline = font["underline"]
            if "color" in font and font["color"]:
                run.font.color.rgb = font["color"]

        # Paragraph overrides
        if "paragraph" in style:
            para = style["paragraph"]
            if "alignment" in para:
                align = para["alignment"].lower()
                if align == "left":
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif align == "center":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == "right":
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif align == "justify":
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
