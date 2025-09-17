# src/templify/core/schema_runner/writers/image_writer.py

from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ImageWriter:
    """
    Handles writing images into the document.
    Schema format:
    {
      "id": "img_1",
      "path": "images/profile.jpg",
      "width": 200,   # pixels
      "height": 200,  # pixels
      "alt_text": "Profile Picture",
      "style": {
        "paragraph": {"alignment": "center"}
      }
    }
    """

    def __init__(self, document):
        self.doc = document

    def write(self, descriptor: dict, style: dict = None):
        """
        Insert an image into the document.
        :param descriptor: schema["images"][i]
        :param style: optional style override dict
        """
        path = descriptor.get("path")
        if not path:
            return None

        width_px = descriptor.get("width")
        height_px = descriptor.get("height")

        # Convert px → Inches (96 dpi assumption, common for Word)
        width = Inches(width_px / 96) if width_px else None
        height = Inches(height_px / 96) if height_px else None

        # Add image
        run = self.doc.add_paragraph().add_run()
        picture = run.add_picture(path, width=width, height=height)

        # Merge style
        merged_style = dict(descriptor.get("style", {}))
        if style:
            merged_style.update(style)

        # Paragraph alignment
        if "paragraph" in merged_style and "alignment" in merged_style["paragraph"]:
            align = merged_style["paragraph"]["alignment"].lower()
            p = run.parent
            if align == "left":
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif align == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == "justify":
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Alt text (python-docx doesn’t support this directly, would need XML injection)
        # We’ll store it in the descriptor for now, for possible XML phase.
        picture._element.get_or_add_cNvPr().descr = descriptor.get("alt_text", "")

        return picture

# ⚠️ Limitations (python-docx)

# Alt text and wrapping/positioning (floating images, inline vs. behind text) require raw XML edits.

# For now, this writer sticks to inline images (which is fine for resumes/reports).