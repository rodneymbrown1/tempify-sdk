# src/templify/core/schema_runner/writers/header_footer_writer.py

from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class HeaderFooterWriter:
    """
    Handles writing headers and footers into the document.
    Schema format:
    {
      "id": "hdr_1",
      "scope": "default",  # default, first_page, even_page
      "location": "header",  # or "footer"
      "style": {
        "style_id": "Header",
        "font": {"name": "Arial", "size": 10, "italic": true},
        "paragraph": {"alignment": "center"}
      },
      "text": "Confidential Resume",
      # or
      "layout": "tabbed",
      "segments": [
        {"text": "[Type here]", "alignment": "left"},
        {"text": "[Type here]", "alignment": "center"},
        {"text": "[Type here]", "alignment": "right"}
      ]
    }
    """

    def __init__(self, document):
        self.doc = document

    def _apply_style(self, p, style: dict):
        """Apply font and paragraph style to the first run/paragraph."""
        if not style:
            return

        # Apply style_id if valid
        if "style_id" in style:
            valid_names = [s.name for s in self.doc.styles]
            if style["style_id"] in valid_names:
                p.style = style["style_id"]

        # Font overrides
        if "font" in style:
            run = p.runs[0] if p.runs else p.add_run()
            font = style["font"]
            if "name" in font:
                run.font.name = font["name"]
            if "size" in font and font["size"]:
                run.font.size = Pt(font["size"])
            if "bold" in font:
                run.font.bold = font["bold"]
            if "italic" in font:
                run.font.italic = font["italic"]
            if "underline" in font:
                run.font.underline = font["underline"]
            if "color" in font and font["color"]:
                try:
                    run.font.color.rgb = RGBColor.from_string(font["color"])
                except Exception:
                    # Fallback: accept raw docx ColorFormat if already set
                    run.font.color.rgb = font["color"]

        # Paragraph properties
        if "paragraph" in style and "alignment" in style["paragraph"]:
            align = style["paragraph"]["alignment"].lower()
            if align == "left":
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif align == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == "justify":
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def write(self, descriptor: dict, style: dict = None):
        """
        Write a header or footer based on schema descriptor.
        :param descriptor: schema["headers"][i] or schema["footers"][i]
        :param style: optional override style dict
        """
        location = descriptor.get("location", "header")  # header or footer
        scope = descriptor.get("scope", "default")

        # For now, apply to the first section (most resumes are single-section docs)
        section = self.doc.sections[0]
        container = section.header if location.lower() == "header" else section.footer

        # Merge style (descriptor + override)
        merged_style = dict(descriptor.get("style", {}))
        if style:
            merged_style.update(style)

        # Handle tabbed layout vs. flat text
        if descriptor.get("layout") == "tabbed" and "segments" in descriptor:
            p = container.add_paragraph()
            for i, seg in enumerate(descriptor["segments"]):
                if i > 0:
                    # Use tab between segments (approximation of tab stops)
                    p.add_run("\t")
                run = p.add_run(seg.get("text", ""))
            self._apply_style(p, merged_style)
        else:
            text = descriptor.get("text", "")
            p = container.add_paragraph(text)
            self._apply_style(p, merged_style)

        return p


# ⚠️ Limitations (python-docx):
# - Cannot set "scope": "first_page"/"even_page" without section property edits.
# - PAGE/NUMPAGES fields still only supported as plain text placeholders ("Page [i]").
# - Tabbed layouts are approximated with literal tabs + default tab stops.
